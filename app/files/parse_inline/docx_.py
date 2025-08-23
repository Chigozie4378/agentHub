from docx import Document

def extract_docx_text(path: str, max_chars: int = 40_000) -> tuple[str, dict]:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
        if sum(len(x) for x in parts) > max_chars:
            break
    return ("\n".join(parts))[:max_chars], {"paragraphs": len(doc.paragraphs)}
